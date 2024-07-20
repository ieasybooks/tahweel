from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tahweel.enums import TransformationType
from tahweel.models import Transformation
from tahweel.utils.string_utils import apply_transformations


NORMALIZE_NUMBERS = str.maketrans('٠١٢٣٤٥٦٧٨٩۰۱۲۳۴۵۶۷۸۹', '01234567890123456789')

TRANSFORMATIONS = [
  Transformation(TransformationType.REPLACE, '\n', ' '),
  Transformation(TransformationType.REPLACE, '\r', ' '),
]


class DocxWriter:
  def __init__(self, file_path: Path):
    self.file_path = file_path

    self.file_path.parent.mkdir(parents=True, exist_ok=True)

  def write(self, texts: list[str], remove_newlines: bool) -> None:
    document = Document()
    allow_rtl = document.styles.add_style('allow_rtl', WD_STYLE_TYPE.CHARACTER)

    for index, text in enumerate(texts):
      paragraph = document.add_paragraph()

      if remove_newlines:
        text = apply_transformations(text, TRANSFORMATIONS)

      run = paragraph.add_run(text.translate(NORMALIZE_NUMBERS))

      paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
      run.style = allow_rtl
      font = run.font
      font.rtl = True

      if index + 1 != len(texts):
        document.add_page_break()

    document.save(self.file_path)
